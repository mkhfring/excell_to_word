import os
import glob

import docx
import click

from excel_to_word.tempelate import SentenceElement, template
from excel_to_word.students import TA


HERE = os.path.dirname(os.path.realpath(__file__))
paragraphs = []


def create_first_paragraph(document, name):
    paragraph1 = f"""
Dear {name},

Thank you very much for applying for our Teaching Assistant (TA) positions. \
We are in the process of finalizing the TA assignments for the first term of the 2021 Winter semester. \
I have the following TA offer for you:  """
    document.add_paragraph(paragraph1)


def create_template(document, element):
    if element[1] == SentenceElement.PARAGRAPH:
        p = document.add_paragraph(element[0])
        paragraphs.append(p)

    if element[1] == SentenceElement.BOLD_PARAGRAPH:
        p = document.add_paragraph()
        p.add_run(element[0]).bold = True
        paragraphs.append(p)

    if element[1] == SentenceElement.SENTENCE:
        p = paragraphs[-1]
        p.add_run(element[0])

    if element[1] == SentenceElement.BOLD:
        p = paragraphs[-1]
        p.add_run(element[0]).bold = True

    if element[1] == SentenceElement.UNDER_LINE:
        p = paragraphs[-1]
        p.add_run(element[0]).underline = True


def add_table(document, student, headers):
    document.add_paragraph(f" {student.duty_hours} hrs/wk for following:")
    assignments_headers = student.assignments[0].__dict__.keys()
    some_table = document.add_table(1, len(headers))
    some_table.style = "Table Grid"
    first_row_cells = some_table.rows[0].cells
    for index, header in enumerate(headers):
        if header in ["TA", "Student No.", "Email"]:
            continue
        first_row_cells[index].text = str(header)

    for index, assignment in enumerate(student.assignments):
        nex_cell = some_table.add_row().cells
        for index, header in enumerate(assignments_headers):
            assert 1 == 1
            nex_cell[index].text = str(getattr(assignment, header))


def create_second_paragraph(document, content):
    for element in content:
        create_template(document, element)


@click.command()
@click.option('--path', default="excell_to_word/data/ta_data.xlsx",
              help='number of greetings')
def main(path):
    print(path)
    # ta = TA("data/ta_data.xlsx")
    ta = TA(os.path.join(path))
    files = glob.glob('data/*')
    for f in files:
        if f.endswith(".xlsx"):
            continue

        os.remove(f)
    for student in ta.students:
        doc = docx.Document()
        create_first_paragraph(doc, student.name)
        add_table(doc, student, ta.data_frame.columns[3:])
        create_second_paragraph(doc, template)
        doc.save(os.path.join(HERE, f"data/{student.name}.docx"))


if __name__ == "__main__":
    main()


