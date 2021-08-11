import os
import glob
import re
import copy
import threading

import docx
import click

from excel_to_word.letter import OfficialLetter, OfferLetter
from excel_to_word.students import TA


HERE = os.path.dirname(os.path.realpath(__file__))
paragraphs = []


def add_table(document, student, headers):
    document.add_paragraph(f" {student.duty_hours} hrs/wk for following:")
    assignments_headers = student.assignments[0].__dict__.keys()
    some_table = document.add_table(1, len(headers))
    some_table.style = "Table Grid"
    first_row_cells = some_table.rows[0].cells
    for index, header in enumerate(headers):
        if header in ["TA", "Student No.", "Email"]:
            continue
        first_row_cells[index].text = str(header)

    for index, assignment in enumerate(student.assignments):
        nex_cell = some_table.add_row().cells
        for index, header in enumerate(assignments_headers):
            assert 1 == 1
            nex_cell[index].text = str(getattr(assignment, header))

    document.add_paragraph()


def handel_paraghraph(document, paragraph):
    para = document.add_paragraph("")
    for run in paragraph.runs:
        output_run = para.add_run(run.text)
        output_run.bold = run.bold
        output_run.italic = run.italic
        output_run.underline = run.underline
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
        # Paragraph's alignment data
    para.paragraph_format.alignment = paragraph.paragraph_format.alignment


def replace_tokens(paraghraph, student):

    new_paragraph = copy.deepcopy(paraghraph)
    name_pattern = r"\[first_name\]"
    if re.search(name_pattern, new_paragraph.text):

        replaced_text = re.sub(name_pattern, student.name.split(" ")[0], new_paragraph.text)
        new_paragraph.text = replaced_text

    return new_paragraph


@click.command()
@click.option('--path', default=os.path.join(HERE, "data/test.xlsx"),
              help='number of greetings')
def main(path):
    files = glob.glob(os.path.join(HERE, 'data/*'))
    for f in files:
        if os.path.isdir(f):
            continue

        if f.endswith(".xlsx"):
            continue

        os.remove(f)

    official_letter = OfficialLetter(
        os.path.join(HERE, "data/test.xlsx"),
        os.path.join(HERE, "templates/letters.docx"),
        os.path.join(HERE, "templates/letters_temp.docx")
    )
    t1 = threading.Thread(target=official_letter.create_output)
    t1.start()
    #official_letter.create_output()
    offer_letter = OfferLetter(
        os.path.join(HERE, "data/test.xlsx"),
        os.path.join(HERE, "templates/offer.docx"),
    )
    t = threading.Thread(target=offer_letter.create_output)
    t.start()
    #offer_letter.create_output()


if __name__ == "__main__":
    main()