import os
import docx


print("Initiating the code")
fack_data = [
    [1, "first record", "second recod"],
    [2, "third record", "fourth recod"]
]
doc = docx.Document()
doc.add_heading("This is a test document", 0)
name = "Mohamad"

body_paragraph1 = f"""
Dear {name},

Thank you very much for applying for our Teaching Assistant (TA) positions.\
 We are in the process of finalizing the TA assignments for the first term of the 2021 Winter semester.\
I have the following TA offer for you:
"""
body_paragraph2 = """
The above hours are an average per week, for 7 weeks (1 week of which is during the final exam period).\
 Hours worked in any one given week may be more or less than the above average.\
  Please note that this is a salaried position and it is the responsibility of the TA and course instructor to manage working hours.\
   If additional hours are required to complete a task, they must be pre-approved by the CMPS Department.\
    Any hours submitted through Workday that was not approved prior will be denied.

"""
bold_scentence1 = "Note that this assignment extends throughout, and possibly beyond, the final exam period"

body_paragraph3 = """ and you may be asked to mark final exams at a date that is """

bold_scentence2 = """after the last day final exams are being written. """
 
body_paragraph4 =  """By accepting this position you acknowledge that you are willing to stay for the full contracted work term, which includes the final exam period.

If you accept this offer, please reply to Chad Davis via e-mail at cdavis.cmpsta@ubc.ca by August 10, 2021. If you can let us know sooner it would be greatly appreciated.

This assignment is not yet finalized and may be subject to change depending on course enrolment figures and TA availability. If necessary you will be contacted regarding any changes, so please keep checking your e-mail.

Kind regards,

Dr. Chad Davis

Lecturer IKBSAS Department of Computer Science, Math, Physics, & Statistics The University of British Columbia | Okanagan Campus
"""
bold_signature = "Dr. Chad Davis"
p1 = doc.add_paragraph(body_paragraph1)

some_table = doc.add_table(1, 3)
some_table.style = "Table Grid"
for id, element1, element2 in fack_data:
    row_cells = some_table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = element1
    row_cells[2].text = element2

p2 = doc.add_paragraph(body_paragraph2)
p2.add_run(bold_scentence1).bold = True
p2.add_run(body_paragraph3)
p2.add_run(bold_scentence2).bold=True
p2.add_run(body_paragraph4)

doc.save("test_doc.docx")
os.system("start test_doc.docx")
