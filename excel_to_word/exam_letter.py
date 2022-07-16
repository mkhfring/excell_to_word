import copy
import glob
import os
import re

import docx
from docx.shared import Pt, RGBColor
import pandas as pd

from excel_to_word.markers import TA


HERE = os.path.dirname(os.path.realpath(__file__))


class CourseData:
    def __init__(self, course, instructor):
        self.course = course
        self.marking_assingments = None
        self.invigilating_assignments = None
        self.instructor = instructor


class ExamLetter:
    def __init__(self, data_path, template_path, output_template_path=None):

        self.table_headers = [
            "Course",
            "Room",
            "Instructor",
            "Instructor Email",
            "Date of exam",
            "time of exam",
            "number of assigned hours"
        ]
        self.teacher_table_headers = [
            "markers",
            "email",
            "ubc email",
            "number of assigned hours"
        ]
        self.output_template_path = output_template_path
        template = docx.Document(template_path)
        template_paragraphs = [paragraph for paragraph in template.paragraphs]
        self.paragraphs = template_paragraphs

    def add_table(self, document, student, role, type='normal'):
        if type == "bullet":

            for index, assinment in enumerate(student.assignments):

                document.add_paragraph(
                    f'\n \t  {str(index + 1)}-  {assinment.subject}  {assinment.course} | {assinment.section} | {assinment.type} | {assinment.days_met} | {assinment.start_time} | {assinment.end_time} | {assinment.hours} |'
                )
            return
        if role == "marking":

            assignments_headers = student.marking_duties[0].__dict__.keys()
            duties = student.marking_duties
        else:
            assignments_headers = student.invigilation_duties[0].__dict__.keys()
            duties = student.invigilation_duties

        some_table = document.add_table(1, len(self.table_headers))
        some_table.style = "Table Grid"
        first_row_cells = some_table.rows[0].cells
        for index, header in enumerate(self.table_headers):
            if header in ["TA", "Student No."]:
                continue

            first_row_cells[index].text = str(header)

        for index, assignment in enumerate(duties):
            nex_cell = some_table.add_row().cells
            for index, header in enumerate(assignments_headers):
                try:
                    nex_cell[index].text = str(getattr(assignment, header))

                except:
                    assert 1 == 1

        document.add_paragraph()

        return self

    def handel_paragraphs(self, document, paragraph, type=0):
        para = document.add_paragraph("")
        for run in paragraph.runs:
            output_run = para.add_run(run.text)
            output_run.bold = run.bold
            output_run.italic = run.italic
            output_run.underline = run.underline
            output_run.font.color.rgb = run.font.color.rgb
            # Run's font data
            output_run.style.name = run.style.name
            # Paragraph's alignment data
        para.paragraph_format.alignment = paragraph.paragraph_format.alignment
        if type:
            style = document.styles['Normal']
            #para.alignment = 0
            font = style.font
            font.name = "Arial"
            font.size = Pt(9)
            para.style = document.styles['Normal']

        return self

    def replace_token(self, paragraph, student):
        NotImplemented

    def create_file(self, student):
        NotImplemented

    def create_output(self):
        NotImplemented


class MarkerLetter(ExamLetter):
    def __init__(self, data_path,  template_path):
        super().__init__(data_path, template_path)
        self.ta_marking = TA(path=data_path, role="student")

    def create_file(self, student):
        doc = docx.Document()
        for paragraph in self.paragraphs:
            if paragraph.text == "":
                continue
            if paragraph.text == "\n":
                continue

            if "table m" in paragraph.text.lower():
                if len(student.marking_duties) > 0:
                    self.add_table(doc, student, 'marking')
                    continue
                else:
                    continue

            if "table i" in paragraph.text.lower():
                if len(student.invigilation_duties) > 0:
                    self.add_table(doc, student, 'invigilation')
                    continue
                else:
                    continue

            paragraph = self.replace_token(paragraph, student)
            self.handel_paragraphs(doc, paragraph, type=1)

        doc.save(os.path.join(HERE, f"data/marker_letters/{student.name}.docx"))
        print(f"The document for {student.name} is saved")

    def replace_token(self, paraghraph, student):
        new_paragraph = copy.deepcopy(paraghraph)
        name_pattern = r"\[first_name\]"
        if re.search(name_pattern, new_paragraph.text):
            replaced_text = re.sub(name_pattern, student.name.split(", ")[1], new_paragraph.text)
            new_paragraph.text = replaced_text

        return new_paragraph

    def create_output(self):
        for student in self.ta_marking.markers:
            self.create_file(student)


class InstructorLetter(ExamLetter):
    def __init__(self, data_path,  template_path):
        super().__init__(data_path, template_path)
        self.teacher_duty = TA(path=data_path, role="teacher")

    def add_table(self, document, course, role, type='normal'):
        if type == "bullet":

            for index, assinment in enumerate(course.assignments):
                document.add_paragraph(
                    f'\n \t  {str(index + 1)}-  {assinment.subject}  {assinment.course} | {assinment.section} | {assinment.type} | {assinment.days_met} | {assinment.start_time} | {assinment.end_time} | {assinment.hours} |'
                )
            return
        if role == "marking":

            assignments_headers = course[0].__dict__.keys()
            duties = course
        else:
            assignments_headers = course[0].__dict__.keys()
            duties = course

        some_table = document.add_table(1, len(self.teacher_table_headers))
        some_table.style = "Table Grid"
        first_row_cells = some_table.rows[0].cells
        for index, header in enumerate(self.teacher_table_headers):
            if header in ["TA", "Student No."]:
                continue

            first_row_cells[index].text = str(header)

        for index, assignment in enumerate(duties):
            nex_cell = some_table.add_row().cells
            for index, header in enumerate(assignments_headers):
                nex_cell[index].text = str(getattr(assignment, header))

        document.add_paragraph()

        return self

    def create_file(self, teacher):
        course_names = []
        for key in teacher.marking_course_names.keys():
            course_names.append(key)
        for key in teacher.invigilating_course_names.keys():
            if key not in course_names:
                course_names.append(key)

        for course in course_names:
            new_data = CourseData(course=course, instructor=teacher.name)
            if course in teacher.marking_course_names.keys():
                new_data.marking_assingments = teacher.marking_course_names[course]
            if course in teacher.invigilating_course_names.keys():
                new_data.invigilating_assignments = teacher.invigilating_course_names[course]

            doc = docx.Document()
            for paragraph in self.paragraphs:
                if paragraph.text == "":
                    continue
                if paragraph.text == "\n":
                    continue

                if "table m" in paragraph.text.lower():
                    if new_data.marking_assingments:
                        self.add_table(doc, new_data.marking_assingments.assignments, 'marking')
                        continue
                    else:
                        continue

                if "table i" in paragraph.text.lower():
                    if new_data.invigilating_assignments:
                        self.add_table(doc, new_data.invigilating_assignments.assignments, 'invigilation')
                        continue
                    else:
                        continue

                paragraph = self.replace_token(paragraph, teacher, course)
                self.handel_paragraphs(doc, paragraph, type=1)

            doc.save(os.path.join(HERE, f"data/instructors_letters/{teacher.name}_{course}.docx"))
            print(f"The document for {teacher.name} is saved")

    def replace_token(self, paraghraph, student, course=None):
        new_paragraph = copy.deepcopy(paraghraph)
        name_pattern = r"\[first_name\]"
        course_pattern = r'\[course\]'
        if re.search(course_pattern, new_paragraph.text):
            replaced_text = re.sub(course_pattern, course, new_paragraph.text)
            new_paragraph.text = replaced_text
            new_paragraph.runs[0].bold = True

        if re.search(name_pattern, new_paragraph.text):
            replaced_text = re.sub(name_pattern, student.name.split(", ")[1], new_paragraph.text)
            new_paragraph.text = replaced_text

        return new_paragraph

    def create_output(self):
        for teacher in self.teacher_duty.teachers:
            self.create_file(teacher)
