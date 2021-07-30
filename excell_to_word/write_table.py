import os
import docx
from excell_to_word.tempelate import SentenceElement, template

paragraphs = []

name = "Mohamad"
paragraph1 = f"""
Dear {name},

Thank you very much for applying for our Teaching Assistant (TA) positions.\
 We are in the process of finalizing the TA assignments for the first term\
  of the 2021 Winter semester. I have the following TA offer for you:
"""


def create_template(document, element):
    if element[1] == SentenceElement.PARAGRAPH:
        p = document.add_paragraph(element[0])
        paragraphs.append(p)

    if element[1] == SentenceElement.BOLD_PARAGRAPH:
        p = document.add_paragraph()
        p.add_run(element[0]).bold = True
        paragraphs.append(p)

    if element[1] == SentenceElement.SENTENCE:
        p = paragraphs[-1]
        p.add_run(element[0])

    if element[1] == SentenceElement.BOLD:
        p = paragraphs[-1]
        p.add_run(element[0]).bold = True

    if element[1] == SentenceElement.UNDER_LINE:
        p = paragraphs[-1]
        p.add_run(element[0]).underline = True


fack_data = [
    [1, "first record", "second recod"],
    [2, "third record", "fourth recod"]
]
doc = docx.Document()
doc.add_paragraph(paragraph1)

some_table = doc.add_table(1, 3)
some_table.style = "Table Grid"
for id, element1, element2 in fack_data:
    row_cells = some_table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = element1
    row_cells[2].text = element2


for element in template:
    create_template(doc, element) 

doc.save("data/test_doc.docx")
os.system("start data/test_doc.docx")
