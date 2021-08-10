from docx import Document
from docx.shared import Pt
import re


def handel_paraghraph(document):
    para = document.add_paragraph()
    for run in paragraph.runs:
        output_run = para.add_run(run.text)
        output_run.bold = run.bold
        output_run.italic = run.italic
        output_run.underline = run.underline
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
        # Paragraph's alignment data

    style = document.styles['Normal']
    para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    para.alignment = 0
    font = style.font
    font.name = "Arial"
    font.size = Pt(9)
    para.stype = document.styles['Normal']



input = Document("templates/letters.docx")
out_put = Document("templates/letters_temp.docx")
footer = out_put.sections[0].footer
for paragraph in footer.paragraphs:
    if re.search("First_Name_", paragraph.text):
        replaced_text = re.sub(
            "Student_ID",
            "123455",
            re.sub(
                "Last_Name_",
                "Khajezade",
                re.sub(
                    "First_Name_",
                    "Mohamad",
                    paragraph.text
                )
            )
        )
        paragraph.text = replaced_text


for paragraph in input.paragraphs:
    # if paragraph.text == "":
    #     continue

    handel_paraghraph(out_put)


# for index, paragraph in enumerate(header.paragraphs):
#     out_put.sections[0].header.paragraphs[index].text = paragraph.text

out_put.save("output.docx")
