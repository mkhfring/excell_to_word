import copy
import glob
import os
import re

import docx
from docx.shared import Pt, RGBColor
import pandas as pd

from excel_to_word.students import TA, required_headers


HERE = os.path.dirname(os.path.realpath(__file__))


class Letter:
    def __init__(self, data_path, template_path, output_template_path=None):
        self.table_headers = [
            "Subject",
            "Course Code",
            "Sec No.",
            "Act Type",
            "Days Met",
            "Start Time",
            "End Time",
            "TA Hours",
        ]
        self.output_template_path = output_template_path
        template = docx.Document(template_path)
        template_paragraphs = [paragraph for paragraph in template.paragraphs]
        self.paragraphs = template_paragraphs
        self.ta_assignment = TA(data_path)

    def add_table(self, document, student, type='normal'):
        if type == "bullet":

            for index, assinment in enumerate(student.assignments):

                document.add_paragraph(
                    f'\n \t  {str(index + 1)}-  {assinment.subject}  {assinment.course} | {assinment.section} | {assinment.type} | {assinment.days_met} | {assinment.start_time} | {assinment.end_time} | {assinment.hours} |'
                )
            return

        document.add_paragraph(f" {student.duty_hours} hrs/wk for following:")
        assignments_headers = student.assignments[0].__dict__.keys()
        some_table = document.add_table(1, len(self.table_headers))
        some_table.style = "Table Grid"
        first_row_cells = some_table.rows[0].cells
        for index, header in enumerate(self.table_headers):
            if header in ["TA", "Student No.", "Email"]:
                continue

            first_row_cells[index].text = str(header)

        for index, assignment in enumerate(student.assignments):
            nex_cell = some_table.add_row().cells
            for index, header in enumerate(assignments_headers):
                nex_cell[index].text = str(getattr(assignment, header))

        document.add_paragraph()
        return self

    def handel_paragraphs(self, document, paragraph, type=0):
        para = document.add_paragraph("")
        for run in paragraph.runs:
            output_run = para.add_run(run.text)
            output_run.bold = run.bold
            output_run.italic = run.italic
            output_run.underline = run.underline
            output_run.font.color.rgb = run.font.color.rgb
            # Run's font data
            output_run.style.name = run.style.name
            # Paragraph's alignment data
        para.paragraph_format.alignment = paragraph.paragraph_format.alignment
        if type:
            style = document.styles['Normal']
            #para.alignment = 0
            font = style.font
            font.name = "Arial"
            font.size = Pt(9)
            para.style = document.styles['Normal']

        return self

    def replace_token(self, paragraph, student):
        NotImplemented

    def create_file(self, student):
        NotImplemented

    def create_output(self):
        for student in self.ta_assignment.students:
            self.create_file(student)


class OfferLetter(Letter):
    def __init__(self, data_path,  template_path):
        super().__init__(data_path, template_path)

    def create_file(self, student):
        doc = docx.Document()
        for paragraph in self.paragraphs:
            if paragraph.text == "":
                continue
            if paragraph.text == "\n":
                continue

            if paragraph.text.lower() == "table":
                self.add_table(doc, student)
                continue

            paragraph = self.replace_token(paragraph, student)
            self.handel_paragraphs(doc, paragraph, type=1)

        doc.save(os.path.join(HERE, f"data/{student.name}.docx"))
        print(f"The document for {student.name} is saved")

    def replace_token(self, paraghraph, student):
        new_paragraph = copy.deepcopy(paraghraph)
        name_pattern = r"\[first_name\]"
        if re.search(name_pattern, new_paragraph.text):
            replaced_text = re.sub(name_pattern, student.name.split(", ")[1], new_paragraph.text)
            new_paragraph.text = replaced_text

        return new_paragraph


class OfficialLetter(Letter):
    def __init__(self, data_path, template_path, output_template_path):
        super().__init__(data_path, template_path, output_template_path)

    def create_file(self, student):

        doc = docx.Document(self.output_template_path)
        self.replace_footer(doc, student)
        for paragraph in self.paragraphs:

            if re.search("\[Table_\]", paragraph.text):
                self.add_table(doc, student, type="bullet")
                continue
            if re.search(r"\[Supervisor_\]", paragraph.text):
                self.add_supervisor(doc, student)
                continue

            paragraph = self.replace_token(paragraph, student)
            self.handel_paragraphs(doc, paragraph, type=1)

        doc.save(os.path.join(HERE, f"data/OfficialLetters/{student.name}.docx"))
        print(f"The document for {student.name} is saved")

    def replace_token(self, paraghraph, student):

        new_paragraph = copy.deepcopy(paraghraph)
        firstname_pattern = r"\[First_Name_\]"
        lastname_pattern = r"\[Last_Name_\]"
        email_pattern = r"\[Email_\]"
        hours_per_week_pattern = r"\[Hours_Per_Week_\]"
        position_pattern = r"\[TA_Position_\]"
        salary_pattern = r"\[Salary_\]"
        total_hours_per_semester_patter = r"\[Total_Hours_In_Semester\]"
        student_name = student.name.split(", ")

        if re.search(firstname_pattern, new_paragraph.text):
            replaced_text = re.sub(firstname_pattern, student.name.split(", ")[1], new_paragraph.text)
            new_paragraph.text = replaced_text

        if re.search(lastname_pattern, new_paragraph.text):
            replaced_text = re.sub(
                lastname_pattern,
                student_name[0] if len(student_name) > 1 else student_name[0],
                new_paragraph.text
            )
            new_paragraph.text = replaced_text

        if re.search(email_pattern, new_paragraph.text):
            replaced_text = re.sub(
                email_pattern,
                student.email,
                new_paragraph.text
            )
            new_paragraph.text = replaced_text

        if re.search(hours_per_week_pattern, new_paragraph.text):
            replaced_text = re.sub(
                hours_per_week_pattern,
                str(student.duty_hours),
                new_paragraph.text
            )
            new_paragraph.text = replaced_text

        if re.search(position_pattern, new_paragraph.text):
            replaced_text = re.sub(
                position_pattern,
                str(student.position) if not pd.isna(student.position) else "{{Not Specified}}",
                new_paragraph.text
            )
            new_paragraph.text = replaced_text

        if re.search(salary_pattern, new_paragraph.text):
            replaced_text = re.sub(
                salary_pattern,
                str(student.salary) if not pd.isna(student.salary) else "{{Not Specified}}",
                new_paragraph.text
            )
            new_paragraph.text = replaced_text

        if re.search(total_hours_per_semester_patter, new_paragraph.text):
            replaced_text = re.sub(
                total_hours_per_semester_patter,
                str(student.total_hours_per_semester),
                new_paragraph.text
            )
            new_paragraph.text = replaced_text

        return new_paragraph

    def replace_footer(self, doc, student):
        footer = doc.sections[0].footer
        name = student.name.split(", ")
        for paragraph in footer.paragraphs:
            if re.search("\[First_Name_\]", paragraph.text):
                replaced_text = re.sub(
                    "\[Student_ID\]",
                    str(int(student.student_id)),
                    re.sub(
                        "\[Last_Name_\]",
                        name[0] if len(name)>1 else name[0],
                        re.sub(
                            "\[First_Name_\]",
                            student.name.split(", ")[1],
                            paragraph.text
                        )
                    )
                )
                paragraph.text = replaced_text

    def add_supervisor(self, doc, student):

        for key, value in student.assigned_courses.items():
            doc.add_paragraph(f"{value[0]} ({value[1]}{key})")

#
# if __name__ == '__main__':
#
#     letter = OfficialLetter(
#         os.path.join(HERE, "data/main_data.xlsx"),
#         os.path.join(HERE, "templates/letters.docx"),
#         os.path.join(HERE, "templates/letters_temp.docx")
#     )
#     letter.create_output()
#     letter = OfferLetter(
#         os.path.join(HERE, "data/main_data.xlsx"),
#         os.path.join(HERE, "templates/offer.docx"),
#     )
#     letter.create_output()
#     assert 1 == 1
